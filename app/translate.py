# translate.py

from __future__ import annotations

import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import TypedDict, List, Dict, Literal, Optional

from docx import Document
from langgraph.graph import StateGraph, END
from openai import OpenAI

from app.chunking import (
    Block,
    BlockType,
    iter_blocks,
    chunk_blocks_with_spans,
    build_contexts,
)
from app.prompts import PROMPTS

# ============================================================
# 0) Types
# ============================================================

Section = Literal["default", "claims", "claims_independent", "claims_dependent", "abstract"]
LayoutMode = Literal["strict", "relaxed", "balanced"]


@dataclass(frozen=True)
class TranslationUnit:
    id: str
    kind: BlockType
    text: str


class TranslateState(TypedDict):
    chunks: List[List[TranslationUnit]]
    contexts: List[Dict[str, str]]
    i: int
    results: Dict[str, str]

    model: str
    base_url: str | None
    api_key: str | None

    target_lang: str

    # routing / prompts
    section: Section
    prompt_default: str
    prompt_claims: str  # legacy/fallback
    prompt_claims_independent: str
    prompt_claims_dependent: str
    prompt_abstract: str

    # abstract word-count tracking
    abstract_word_count: int
    last_abstract_block_id: str | None
    abstract_completed: bool  # marks that we've seen abstract body text
    abstract_block_ids: List[str]
    abstract_target_words: int
    abstract_word_tolerance: int
    abstract_rewritten: bool

    # cross-chunk context for terminology consistency
    glossary: Dict[str, str]       # Korean term → English term (accumulated)
    prev_translated_text: str      # English output of the previous chunk

    # cross-chunk context for claim preambles
    claim_preambles: Dict[str, str]  # claim number → category (e.g. "1" → "method")


# ============================================================
# 1) OpenAI client + JSON extraction fallback
# ============================================================

def make_client(base_url: Optional[str] = None, api_key: Optional[str] = None) -> OpenAI:
    kwargs = {"timeout": 120.0, "max_retries": 2}
    if base_url:
        kwargs["base_url"] = base_url
    if api_key:
        kwargs["api_key"] = api_key
    return OpenAI(**kwargs)


def extract_first_json_object(s: str) -> Optional[str]:
    s2 = (s or "").strip()
    if s2.startswith("{") and s2.endswith("}"):
        return s2

    start = s2.find("{")
    if start == -1:
        return None

    depth = 0
    for i in range(start, len(s2)):
        ch = s2[i]
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return s2[start : i + 1]
    return None


# ============================================================
# 2) Translate one chunk with a given prompt
# ============================================================

def _format_glossary(glossary: Dict[str, str]) -> str:
    if not glossary:
        return "(none yet — this is the first chunk)"
    lines = [f"{ko} → {en}" for ko, en in glossary.items()]
    return "\n".join(lines)


def _format_claim_preambles(preambles: Dict[str, str]) -> str:
    if not preambles:
        return "(none yet — no independent claims translated so far)"
    lines = [
        f"Claim {num}: {cat}"
        for num, cat in sorted(
            preambles.items(),
            key=lambda x: int(x[0]) if x[0].isdigit() else 0
        )
    ]
    return "\n".join(lines)


_INDEP_CLAIM_RE = re.compile(r"^\s*(\d+)\.\s+An?\s+(\w+)", re.MULTILINE)

_CATEGORY_WORDS = {
    "method", "apparatus", "system", "device", "medium", "program",
    "composition", "compound", "kit", "assembly", "machine", "article",
    "process", "circuit", "network", "computer", "storage", "structure",
    "sensor", "module", "unit", "arrangement",
}


def extract_claim_preambles(translated_text: str) -> Dict[str, str]:
    """Regex fallback: extract independent claim preambles from translated English text."""
    preambles: Dict[str, str] = {}
    for m in _INDEP_CLAIM_RE.finditer(translated_text):
        claim_num = m.group(1)
        category = m.group(2).lower()
        if category in _CATEGORY_WORDS:
            preambles[claim_num] = category
    return preambles


def translate_chunk(
    client: OpenAI,
    model: str,
    *,
    prompt_name: str,
    target_lang: str,
    chunk: List[TranslationUnit],
    context_before: str = "",
    context_after: str = "",
    glossary: Dict[str, str] | None = None,
    prev_translated_text: str = "",
    claim_preambles: Dict[str, str] | None = None,
    abstract_target_words: int = 0,
) -> tuple[Dict[str, str], List[Dict[str, str]], Dict[str, str]]:
    """Translate a chunk and return (translations, key_terms, new_claim_preambles)."""
    prompt = PROMPTS[prompt_name]
    payload = [{"id": b.id, "text": b.text} for b in chunk]

    glossary_str = _format_glossary(glossary or {})
    prev_translation = prev_translated_text or "(none — this is the first chunk)"
    preambles_str = _format_claim_preambles(claim_preambles or {})

    system_prompt = prompt.system
    user_prompt = prompt.render_user(
        payload=payload,
        target_lang=target_lang,
        context_before=context_before,
        context_after=context_after,
        glossary=glossary_str,
        prev_translation=prev_translation,
        claim_preambles=preambles_str,
    )
    if abstract_target_words > 0 and "abstract" in prompt_name:
        user_prompt += (
            "\nADDITIONAL ABSTRACT CONSTRAINT:\n"
            f"- Target exactly {abstract_target_words} English words for the abstract body text.\n"
            "- Do not include headings in the count.\n"
        )

    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.0,
    )

    content = (resp.choices[0].message.content or "").strip()

    try:
        data = json.loads(content)
    except Exception:
        json_str = extract_first_json_object(content)
        if not json_str:
            raise ValueError(f"Model did not return JSON. Got:\n{content[:800]}")
        data = json.loads(json_str)

    out: Dict[str, str] = {}
    translations = data.get("translations", [])
    if isinstance(translations, list):
        for item in translations:
            if isinstance(item, dict) and "id" in item and "text" in item:
                out[str(item["id"])] = str(item["text"])

    # Extract key_terms for glossary accumulation
    key_terms: List[Dict[str, str]] = []
    raw_terms = data.get("key_terms", [])
    if isinstance(raw_terms, list):
        for term in raw_terms:
            if isinstance(term, dict) and "ko" in term and "en" in term:
                key_terms.append({"ko": str(term["ko"]), "en": str(term["en"])})

    # Extract claim_preambles from LLM response (primary)
    new_preambles: Dict[str, str] = {}
    raw_preambles = data.get("claim_preambles", [])
    if isinstance(raw_preambles, list):
        for p in raw_preambles:
            if isinstance(p, dict) and "claim_num" in p and "category" in p:
                new_preambles[str(p["claim_num"])] = str(p["category"]).lower()

    # Regex fallback: extract from translated text if LLM didn't provide them
    if not new_preambles and out:
        translated_text = "\n".join(out.values())
        new_preambles = extract_claim_preambles(translated_text)

    return out, key_terms, new_preambles


# ============================================================
# 3) Abstract word counting + finalization
# ============================================================

_WORD_RE = re.compile(r"[A-Za-z]+(?:'[A-Za-z]+)?")

def count_english_words(text: str) -> int:
    return len(_WORD_RE.findall(text or ""))

def append_word_count_to_last_sentence(text: str, n: int) -> str:
    s = (text or "").rstrip()
    if not s:
        return s
    if re.search(r"\(\d+\s+words\)\s*$", s):
        return s
    return f"{s} ({n} words)"


def _strip_abstract_heading(text: str) -> str:
    return re.sub(r"^ABSTRACT[:\s\-]*", "", (text or "").strip(), flags=re.IGNORECASE).strip()


def _call_abstract_rewrite(
    *,
    client: OpenAI,
    model: str,
    abstract_text: str,
    target_words: int,
) -> str:
    system_prompt = (
        "You are a patent abstract editor. Rewrite for legal-faithful clarity.\n"
        "Return only the rewritten abstract text without markdown, labels, or explanations."
    )
    user_prompt = (
        f"Rewrite the following abstract to be exactly {target_words} English words.\n"
        "Keep technical meaning and scope unchanged.\n"
        "Do not add numbering, headings, or commentary.\n\n"
        f"ABSTRACT:\n{abstract_text}"
    )
    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.0,
    )
    out = (resp.choices[0].message.content or "").strip()
    out = re.sub(r"^```(?:text)?\s*", "", out, flags=re.IGNORECASE).strip()
    out = re.sub(r"\s*```$", "", out).strip()
    return _strip_abstract_heading(out)


# ============================================================
# 4) Section detection + routing
# ============================================================

def _chunk_text(chunk: List[TranslationUnit]) -> str:
    return "\n".join((b.text or "") for b in chunk)

def detect_section_from_chunk(prev: Section, chunk: List[TranslationUnit]) -> Section:
    """
    Claims starts with 【청구 범위】 (sometimes without space).
    Abstract starts with 【요약】, but many KR docs also have 【요약서】.
    Sticky behavior: if no heading is found, keep previous section.
    """
    t = _chunk_text(chunk)

    # If we're already in abstract, keep it sticky (caller decides exit timing)
    if prev == "abstract":
        return "abstract"

    # Claims
    if "【청구" in t:
        return "claims"

    # Abstract: match 【요약】 and 【요약서】 and similar forms
    if "【요약" in t:
        return "abstract"

    return prev


_DEP_CLAIM_RE = re.compile(r"(청구항\s*\d+|제\s*\d+\s*항)\s*(에\s*따른|에\s*따라|에\s*있어서|의)")
# Abstract-ending / non-abstract markers (대표도 + FIG lines)
_FIGURE_RE = re.compile(
    r"(【\s*대표도\s*】|대표도|도\s*\d+[A-Za-z]?|FIG\.\s*\d+|도면의\s*간단한\s*설명)"
)

def is_dependent_claim_chunk(chunk: List[TranslationUnit]) -> bool:
    t = _chunk_text(chunk)
    return bool(_DEP_CLAIM_RE.search(t))

def is_figure_description_chunk(chunk: List[TranslationUnit]) -> bool:
    """
    Heuristic: 대표도 / 도 1 / FIG. 1 / brief description of drawings often follow the abstract.
    If detected, we should exit the abstract section to avoid counting their words.
    """
    t = _chunk_text(chunk)
    return bool(_FIGURE_RE.search(t))


def node_route_section(state: TranslateState) -> TranslateState:
    i = state["i"]

    def finalize_abstract_word_count(s: TranslateState) -> TranslateState:
        n = int(s.get("abstract_word_count", 0) or 0)
        last_id = s.get("last_abstract_block_id")
        results = s.get("results", {})
        abstract_ids = list(s.get("abstract_block_ids", []))
        target = int(s.get("abstract_target_words", 0) or 0)
        tol = int(s.get("abstract_word_tolerance", 0) or 0)
        rewritten = bool(s.get("abstract_rewritten", False))

        new_results = dict(results)

        # Enforce target words with a rewrite pass when configured.
        if (
            target > 0
            and not rewritten
            and abstract_ids
            and abs(n - target) > tol
        ):
            combined = "\n".join(
                _strip_abstract_heading(new_results.get(bid, ""))
                for bid in abstract_ids
                if (new_results.get(bid, "") or "").strip()
            ).strip()
            if combined:
                try:
                    client = make_client(s.get("base_url"), s.get("api_key"))
                    rewritten_text = _call_abstract_rewrite(
                        client=client,
                        model=s["model"],
                        abstract_text=combined,
                        target_words=target,
                    )
                    if rewritten_text:
                        # Keep abstract as a single coherent paragraph on the last abstract id.
                        for bid in abstract_ids:
                            if bid != last_id:
                                new_results[bid] = ""
                        if last_id:
                            new_results[last_id] = rewritten_text
                        n = count_english_words(rewritten_text)
                        rewritten = True
                        print(f"[ABSTRACT] rewrite applied: {n} words (target={target}, tol={tol})")
                except Exception as e:
                    print(f"[ABSTRACT] rewrite skipped due to error: {e}")

        if n > 0 and last_id and last_id in new_results:
            new_results[last_id] = append_word_count_to_last_sentence(new_results[last_id], n)
            return {
                **s,
                "results": new_results,
                "abstract_completed": True,
                "abstract_word_count": n,
                "abstract_rewritten": rewritten,
            }

        return {**s, "results": new_results, "abstract_rewritten": rewritten}

    # If we are done, finalize abstract once here
    if i >= len(state["chunks"]):
        if state.get("section") == "abstract" or state.get("abstract_word_count", 0):
            return finalize_abstract_word_count(state)
        return state

    prev = state.get("section", "default")
    new = detect_section_from_chunk(prev, state["chunks"][i])

    # If we're in abstract and see representative-figure / drawing markers, exit abstract
    if prev == "abstract" and is_figure_description_chunk(state["chunks"][i]):
        new = "default"

    # Leaving abstract -> finalize "(XXX words)" on last abstract sentence
    if prev == "abstract" and new != "abstract":
        finalized = finalize_abstract_word_count(state)
        if finalized is not state:
            print(f"[ABSTRACT] finalize word count: {finalized.get('abstract_word_count', 0)} words")
            state = finalized

    if new != prev:
        print(f"[ROUTE] chunk {i}: {prev} -> {new}")

    return {**state, "section": new}


def route_after_router(state: TranslateState) -> str:
    if state["i"] >= len(state["chunks"]):
        return END

    sec = state.get("section", "default")
    if sec == "claims":
        chunk = state["chunks"][state["i"]]
        if is_dependent_claim_chunk(chunk):
            return "translate_claims_dependent"
        return "translate_claims_independent"
    if sec == "claims_independent":
        return "translate_claims_independent"
    if sec == "claims_dependent":
        return "translate_claims_dependent"
    if sec == "abstract":
        return "translate_abstract"
    return "translate_default"


# ============================================================
# 5) Translation nodes (default/claims/abstract)
# ============================================================

def _translate_with_prompt(state: TranslateState, prompt_name: str) -> TranslateState:
    i = state["i"]
    chunks = state["chunks"]
    contexts = state.get("contexts", [])
    if i >= len(chunks):
        return state

    client = make_client(state.get("base_url"), state.get("api_key"))
    chunk = chunks[i]
    ctx_before = contexts[i].get("before", "") if i < len(contexts) else ""
    ctx_after = contexts[i].get("after", "") if i < len(contexts) else ""
    sec = state.get("section", "default")
    glossary = dict(state.get("glossary", {}))
    prev_translated_text = state.get("prev_translated_text", "")
    claim_preambles = dict(state.get("claim_preambles", {}))
    abstract_target_words = int(state.get("abstract_target_words", 0) or 0)

    print(f"[{sec}] {i}-th chunk... (glossary: {len(glossary)} terms, preambles: {len(claim_preambles)})")

    translated_map, key_terms, new_preambles = translate_chunk(
        client=client,
        model=state["model"],
        prompt_name=prompt_name,
        target_lang=state["target_lang"],
        chunk=chunk,
        context_before=ctx_before,
        context_after=ctx_after,
        glossary=glossary,
        prev_translated_text=prev_translated_text,
        claim_preambles=claim_preambles,
        abstract_target_words=abstract_target_words,
    )

    results = dict(state["results"])
    results.update(translated_map)

    # Accumulate glossary with newly extracted terms
    for term in key_terms:
        glossary[term["ko"]] = term["en"]

    # Accumulate claim preambles
    claim_preambles.update(new_preambles)

    # Build prev_translated_text from this chunk's English output
    new_prev = "\n".join(translated_map.get(b.id, b.text) for b in chunk)

    # Track abstract word count
    abstract_word_count = int(state.get("abstract_word_count", 0))
    last_abstract_block_id = state.get("last_abstract_block_id")
    abstract_completed = state.get("abstract_completed", False)
    abstract_block_ids = list(state.get("abstract_block_ids", []))

    if sec == "abstract":
        for b in chunk:
            bid = b.id
            txt = translated_map.get(bid, "")
            t = (txt or "").strip()

            # Remove ABSTRACT heading if model attaches it
            t_no_heading = re.sub(r"^ABSTRACT[:\s\-]*", "", t, flags=re.IGNORECASE).strip()

            # Skip pure heading lines
            if not t_no_heading or t_no_heading.upper() == "ABSTRACT":
                continue

            # Skip 대표도/FIG captions from word count (common end-of-doc pattern)
            if re.match(r"^(REPRESENTATIVE\s+FIGURE)\b", t_no_heading, flags=re.IGNORECASE):
                continue
            if re.match(r"^(FIG\.)\s*\d+", t_no_heading, flags=re.IGNORECASE):
                continue

            wc = count_english_words(t_no_heading)
            if wc > 0:
                abstract_word_count += wc
                last_abstract_block_id = bid
                if bid not in abstract_block_ids:
                    abstract_block_ids.append(bid)

        if abstract_word_count > 0:
            abstract_completed = True

    return {
        **state,
        "results": results,
        "i": i + 1,
        "abstract_word_count": abstract_word_count,
        "last_abstract_block_id": last_abstract_block_id,
        "abstract_completed": abstract_completed,
        "abstract_block_ids": abstract_block_ids,
        "glossary": glossary,
        "prev_translated_text": new_prev,
        "claim_preambles": claim_preambles,
    }


def node_translate_default(state: TranslateState) -> TranslateState:
    return _translate_with_prompt(state, state["prompt_default"])

def node_translate_claims(state: TranslateState) -> TranslateState:
    return _translate_with_prompt(state, state["prompt_claims"])

def node_translate_claims_independent(state: TranslateState) -> TranslateState:
    return _translate_with_prompt(state, state["prompt_claims_independent"])

def node_translate_claims_dependent(state: TranslateState) -> TranslateState:
    return _translate_with_prompt(state, state["prompt_claims_dependent"])

def node_translate_abstract(state: TranslateState) -> TranslateState:
    return _translate_with_prompt(state, state["prompt_abstract"])


# ============================================================
# 6) Build graph
# ============================================================

def build_translation_graph():
    g = StateGraph(TranslateState)

    g.add_node("route_section", node_route_section)
    g.add_node("translate_default", node_translate_default)
    g.add_node("translate_claims", node_translate_claims)
    g.add_node("translate_claims_independent", node_translate_claims_independent)
    g.add_node("translate_claims_dependent", node_translate_claims_dependent)
    g.add_node("translate_abstract", node_translate_abstract)

    g.set_entry_point("route_section")

    g.add_conditional_edges(
        "route_section",
        route_after_router,
        {
            "translate_default": "translate_default",
            "translate_claims": "translate_claims",
            "translate_claims_independent": "translate_claims_independent",
            "translate_claims_dependent": "translate_claims_dependent",
            "translate_abstract": "translate_abstract",
            END: END,
        },
    )

    g.add_edge("translate_default", "route_section")
    g.add_edge("translate_claims", "route_section")
    g.add_edge("translate_claims_independent", "route_section")
    g.add_edge("translate_claims_dependent", "route_section")
    g.add_edge("translate_abstract", "route_section")

    return g.compile()


# ============================================================
# 7) Translation units + writeback mapping
# ============================================================

_HEADING_RE = re.compile(r"^\s*【[^】]+】\s*$")
_CLAIM_NUMBER_LINE_RE = re.compile(r"^\s*\d+\s*[\.\)]")
_LIST_PREFIX_RE = re.compile(r"^\s*[-*•]\s+")
_HARD_SENTENCE_END_RE = re.compile(r"[.!?;:]\s*$")
_SENTENCE_SPLIT_RE = re.compile(r"(?<=[.!?;:])\s+")


def _same_table_cell(a: Block, b: Block) -> bool:
    if not (a.id.startswith("cell:") and b.id.startswith("cell:")):
        return False
    a_parts = a.id.split(":")
    b_parts = b.id.split(":")
    if len(a_parts) != 5 or len(b_parts) != 5:
        return False
    return a_parts[:4] == b_parts[:4]


def _can_merge_blocks(left: Block, right: Block) -> bool:
    lt = (left.text or "").strip()
    rt = (right.text or "").strip()
    if not lt or not rt:
        return False
    if left.kind != right.kind:
        return False
    if left.kind == "cell_p" and not _same_table_cell(left, right):
        return False
    if _HEADING_RE.match(lt) or _HEADING_RE.match(rt):
        return False
    if _CLAIM_NUMBER_LINE_RE.match(lt) or _CLAIM_NUMBER_LINE_RE.match(rt):
        return False
    if _LIST_PREFIX_RE.match(rt):
        return False
    if _HARD_SENTENCE_END_RE.search(lt):
        return False

    # Conservative merge for wrapped lines: keep merges short.
    return len(lt) <= 160 and len(rt) <= 160


def build_translation_units(
    blocks: List[Block],
    *,
    layout_mode: LayoutMode = "strict",
) -> tuple[List[TranslationUnit], Dict[str, List[str]]]:
    if layout_mode == "strict":
        units = [TranslationUnit(id=b.id, kind=b.kind, text=b.text) for b in blocks]
        return units, {b.id: [b.id] for b in blocks}

    units: List[TranslationUnit] = []
    mapping: Dict[str, List[str]] = {}
    i = 0
    n = len(blocks)
    uidx = 0

    while i < n:
        first = blocks[i]
        member_ids = [first.id]
        member_texts = [first.text or ""]
        j = i + 1

        if j < n and _can_merge_blocks(first, blocks[j]):
            member_ids.append(blocks[j].id)
            member_texts.append(blocks[j].text or "")
            j += 1

        unit_id = f"u:{uidx}"
        units.append(TranslationUnit(id=unit_id, kind=first.kind, text="\n".join(member_texts)))
        mapping[unit_id] = member_ids

        uidx += 1
        i = j

    return units, mapping


def _balanced_split_text(text: str, n_parts: int) -> List[str]:
    t = (text or "").strip()
    if n_parts <= 1:
        return [t]
    if not t:
        return [""] * n_parts

    sents = [s.strip() for s in _SENTENCE_SPLIT_RE.split(t) if s.strip()]
    if len(sents) >= n_parts:
        out = ["" for _ in range(n_parts)]
        for idx, s in enumerate(sents):
            pos = min(idx * n_parts // len(sents), n_parts - 1)
            out[pos] = (out[pos] + " " + s).strip()
        return out

    words = t.split()
    if not words:
        return [""] * n_parts
    out = ["" for _ in range(n_parts)]
    for idx, w in enumerate(words):
        pos = min(idx * n_parts // len(words), n_parts - 1)
        out[pos] = (out[pos] + " " + w).strip()
    return out


def expand_unit_translations(
    unit_translations: Dict[str, str],
    unit_to_block_ids: Dict[str, List[str]],
    *,
    layout_mode: LayoutMode,
) -> Dict[str, str]:
    out: Dict[str, str] = {}
    for uid, translated in unit_translations.items():
        block_ids = unit_to_block_ids.get(uid, [])
        if not block_ids:
            continue

        if len(block_ids) == 1 or layout_mode == "strict":
            out[block_ids[0]] = translated
            for bid in block_ids[1:]:
                out[bid] = ""
            continue

        if layout_mode == "balanced":
            parts = _balanced_split_text(translated, len(block_ids))
            for bid, part in zip(block_ids, parts):
                out[bid] = part
            continue

        # relaxed: place reconstructed sentence on first line, clear follower lines
        out[block_ids[0]] = translated
        for bid in block_ids[1:]:
            out[bid] = ""
    return out


# ============================================================
# 8) Apply translations back into DOCX
# ============================================================

def apply_translations_to_docx(
    src_docx_path: str | Path,
    translations: Dict[str, str],
    out_docx_path: str | Path,
) -> None:
    doc = Document(str(src_docx_path))

    for block_id, translated in translations.items():
        if block_id.startswith("p:"):
            _, idx = block_id.split(":", 1)
            pi = int(idx)
            if 0 <= pi < len(doc.paragraphs):
                doc.paragraphs[pi].text = translated

        elif block_id.startswith("cell:"):
            parts = block_id.split(":")
            if len(parts) == 5:
                _, ti, ri, ci, pi = parts
                ti = int(ti); ri = int(ri); ci = int(ci); pi = int(pi)

                if 0 <= ti < len(doc.tables):
                    table = doc.tables[ti]
                    if 0 <= ri < len(table.rows):
                        row = table.rows[ri]
                        if 0 <= ci < len(row.cells):
                            cell = row.cells[ci]
                            if 0 <= pi < len(cell.paragraphs):
                                cell.paragraphs[pi].text = translated

    doc.save(str(out_docx_path))


# ============================================================
# 9) End-to-end runner
# ============================================================

def translate_docx(
    src_docx_path: str | Path,
    out_docx_path: str | Path,
    *,
    model: str,
    target_lang: str = "English",
    max_chars_per_chunk: int = 2500,
    context_chars: int = 1000,
    chunk_overlap: int = 0,
    layout_mode: LayoutMode = "relaxed",
    abstract_target_words: int = 0,
    abstract_word_tolerance: int = 5,
    base_url: Optional[str] = None,
    api_key: Optional[str] = None,
    prompt_default: str = "patent_kr2en_body_v1",
    prompt_claims: str = "patent_kr2en_claims_v1",
    prompt_claims_independent: str = "patent_kr2en_claims_independent_v1",
    prompt_claims_dependent: str = "patent_kr2en_claims_dependent_v1",
    prompt_abstract: str = "patent_kr2en_abstract_v1",
) -> None:
    print("Start chunking...")
    blocks = iter_blocks(src_docx_path)
    units, unit_to_block_ids = build_translation_units(blocks, layout_mode=layout_mode)
    chunk_spans = chunk_blocks_with_spans(
        units,
        max_chars=max_chars_per_chunk,
        overlap=chunk_overlap,
    )
    chunks = [c for (c, _, _) in chunk_spans]
    contexts_raw = build_contexts(
        units,
        [(s, e) for (_, s, e) in chunk_spans],
        context_chars=context_chars,
    )
    contexts = [{"before": b, "after": a} for (b, a) in contexts_raw]
    print(
        f"Finish chunking: {len(chunks)} chunks from {len(units)} units / {len(blocks)} blocks "
        f"(context ±{context_chars} chars, layout_mode={layout_mode})"
    )

    print("Build graph...")
    app = build_translation_graph()

    print("Start invoking...")
    final = app.invoke(
        {
            "chunks": chunks,
            "contexts": contexts,
            "i": 0,
            "results": {},
            "model": model,
            "base_url": base_url,
            "api_key": api_key,
            "target_lang": target_lang,

            "section": "default",
            "prompt_default": prompt_default,
            "prompt_claims": prompt_claims,
            "prompt_claims_independent": prompt_claims_independent,
            "prompt_claims_dependent": prompt_claims_dependent,
            "prompt_abstract": prompt_abstract,

            "abstract_word_count": 0,
            "last_abstract_block_id": None,
            "abstract_completed": False,
            "abstract_block_ids": [],
            "abstract_target_words": abstract_target_words,
            "abstract_word_tolerance": abstract_word_tolerance,
            "abstract_rewritten": False,

            "glossary": {},
            "prev_translated_text": "",
            "claim_preambles": {},
        },
        config={"recursion_limit": max(50, len(chunks) * 3)},
    )

    # Defensive finalization: if document ends while in abstract (or word count exists), finalize
    if final.get("abstract_word_count", 0) or final.get("section") == "abstract":
        final = node_route_section({**final, "i": len(chunks)})

    print("Apply translations...")
    block_translations = expand_unit_translations(
        final["results"],
        unit_to_block_ids,
        layout_mode=layout_mode,
    )
    apply_translations_to_docx(src_docx_path, block_translations, out_docx_path)
    print("Done.")
