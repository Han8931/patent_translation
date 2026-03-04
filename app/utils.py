# utils.py

import os
import boto3
import re

from typing import Optional
from docx import Document
from pathlib import Path
from fastapi import HTTPException  # (currently unused)

from botocore.exceptions import BotoCoreError, ClientError  # (currently unused)

S3_DRIVE = "http://s3.dataplatform.samsungds.net:9020"
FILE_DOWNLOAD_DIR = Path("./uploads/")


def extract_docx_text(docx_path: str | Path) -> str:
    docx_path = Path(docx_path)
    doc = Document(str(docx_path))

    parts: list[str] = []

    # Paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    # Tables (optional but usually important)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            line = "\t".join([c for c in cells if c])
            if line:
                parts.append(line)

    return "\n".join(parts)


def load_filename_list(file_path: Path) -> list[str]:
    lines = file_path.read_text(encoding="utf-8").splitlines()
    out: list[str] = []

    for line in lines:
        s = line.strip()
        if not s or s.startswith("#"):
            continue

        # remove surrounding quotes if any
        if (s.startswith('"') and s.endswith('"')) or (s.startswith("'") and s.endswith("'")):
            s = s[1:-1].strip()

        # if the line is a path, keep only the filename
        out.append(Path(s).name)

    return out


def download_s3_file(key: str | Path, bucket: "Bucket") -> Path:
    key = str(key)  # IMPORTANT: boto3 Key must be str

    client = boto3.client(
        "s3",
        aws_access_key_id=str(bucket.access_key.get_secret_value()),
        aws_secret_access_key=str(bucket.secret_key.get_secret_value()),
        endpoint_url=S3_DRIVE,
    )

    filename = Path(key).name
    path = FILE_DOWNLOAD_DIR / filename

    os.makedirs(FILE_DOWNLOAD_DIR, exist_ok=True)

    client.download_file(
        Bucket=bucket.name,
        Key=key,                 # now guaranteed str
        Filename=str(path),
    )
    return path


def list_s3_keys(
    bucket: "Bucket",
    prefix: str,
    *,
    endpoint_url: str = S3_DRIVE,
    suffix: Optional[str] = ".docx",
    include_regex: Optional[str] = None,
    exclude_regex: Optional[str] = None,
) -> list[str]:
    """
    List object keys in S3 under `prefix`.
    Returns keys like: 'patent/foo.docx'
    """
    client = boto3.client(
        "s3",
        aws_access_key_id=str(bucket.access_key.get_secret_value()),
        aws_secret_access_key=str(bucket.secret_key.get_secret_value()),
        endpoint_url=endpoint_url,
    )

    include_re = re.compile(include_regex) if include_regex else None
    exclude_re = re.compile(exclude_regex) if exclude_regex else None

    keys: list[str] = []
    token = None

    while True:
        kwargs = {"Bucket": bucket.name, "Prefix": prefix, "MaxKeys": 1000}
        if token:
            kwargs["ContinuationToken"] = token

        resp = client.list_objects_v2(**kwargs)

        for obj in resp.get("Contents", []):
            key = obj["Key"]

            # skip "folders"
            if key.endswith("/"):
                continue

            if suffix and not key.lower().endswith(suffix.lower()):
                continue

            if include_re and not include_re.search(key):
                continue

            if exclude_re and exclude_re.search(key):
                continue

            keys.append(key)

        if resp.get("IsTruncated"):
            token = resp.get("NextContinuationToken")
        else:
            break

    return keys
