from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Literal, Sequence, Tuple

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P


# Basic block unit used throughout translation.
BlockType = Literal["p", "cell_p"]


@dataclass
class Block:
    id: str
    kind: BlockType
    text: str


def iter_blocks(src_docx_path: str | Path) -> List[Block]:
    """
    Walk the DOCX in document order and emit paragraph and table-cell paragraphs as Block objects.
    The ids follow the scheme consumed by apply_translations_to_docx:
      - Paragraph: p:{para_index}
      - Table cell paragraph: cell:{table_idx}:{row_idx}:{col_idx}:{para_index_in_cell}
    """
    doc = Document(str(src_docx_path))
    blocks: List[Block] = []

    # Map table XML element to its table index for stable ids
    table_elements = {tbl._element: i for i, tbl in enumerate(doc.tables)}

    para_counter = 0  # matches doc.paragraphs indices (table paragraphs excluded)

    for child in doc.element.body:
        if isinstance(child, CT_P):
            para = doc.paragraphs[para_counter]
            blocks.append(Block(id=f"p:{para_counter}", kind="p", text=para.text or ""))
            para_counter += 1
        elif isinstance(child, CT_Tbl):
            ti = table_elements.get(child, None)
            if ti is None or ti >= len(doc.tables):
                continue
            table = doc.tables[ti]
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    for pi, para in enumerate(cell.paragraphs):
                        blocks.append(
                            Block(
                                id=f"cell:{ti}:{ri}:{ci}:{pi}",
                                kind="cell_p",
                                text=para.text or "",
                            )
                        )

    return blocks


def _block_length(b: Block) -> int:
    # Add 1 char to account for a joining newline when chunking.
    return len(b.text or "") + 1


def chunk_blocks(
    blocks: Sequence[Block] | Iterable[Block],
    *,
    max_chars: int = 2500,
    overlap: int = 0,
) -> List[List[Block]]:
    """
    Group sequential blocks into chunks that respect max_chars.
    Optional block-level overlap allows a bit of preceding context to travel with
    the next chunk (without re-translating the overlapping blocks).
    """
    if not isinstance(blocks, Sequence):
        blocks = list(blocks)

    chunks: List[List[Block]] = []
    start = 0
    n = len(blocks)

    while start < n:
        chunk: List[Block] = []
        total = 0
        idx = start

        while idx < n:
            b = blocks[idx]
            b_len = _block_length(b)

            # If adding this block would exceed the limit, finish current chunk.
            if chunk and total + b_len > max_chars:
                break

            chunk.append(b)
            total += b_len
            idx += 1

        # If a single block is larger than max_chars, force it alone.
        if not chunk:
            chunk = [blocks[idx]]
            idx += 1

        chunks.append(chunk)

        # Advance start pointer with optional overlap
        next_start = idx - overlap
        if next_start <= start:
            next_start = idx  # ensure progress
        start = next_start

    return chunks


def build_contexts(
    blocks: Sequence[Block],
    spans: Sequence[Tuple[int, int]],
    *,
    context_chars: int = 400,
) -> List[Tuple[str, str]]:
    """
    For each (start, end) span (end exclusive), build surrounding context strings.
    Returns list of (context_before, context_after).
    """
    def gather_before(start: int) -> str:
        out: List[str] = []
        total = 0
        idx = start - 1
        while idx >= 0 and total < context_chars:
            t = blocks[idx].text or ""
            total += len(t) + 1
            out.append(t)
            idx -= 1
        return "\n".join(reversed(out))

    def gather_after(end: int) -> str:
        out: List[str] = []
        total = 0
        idx = end
        n = len(blocks)
        while idx < n and total < context_chars:
            t = blocks[idx].text or ""
            total += len(t) + 1
            out.append(t)
            idx += 1
        return "\n".join(out)

    return [(gather_before(s), gather_after(e)) for s, e in spans]


def chunk_blocks_with_spans(
    blocks: Sequence[Block] | Iterable[Block],
    *,
    max_chars: int = 2500,
    overlap: int = 0,
) -> List[Tuple[List[Block], int, int]]:
    """
    Variant of chunk_blocks that also returns (start_idx, end_idx) spans.
    """
    if not isinstance(blocks, Sequence):
        blocks = list(blocks)

    chunks: List[Tuple[List[Block], int, int]] = []
    start = 0
    n = len(blocks)

    while start < n:
        chunk: List[Block] = []
        total = 0
        idx = start

        while idx < n:
            b = blocks[idx]
            b_len = _block_length(b)
            if chunk and total + b_len > max_chars:
                break
            chunk.append(b)
            total += b_len
            idx += 1

        if not chunk:
            chunk = [blocks[idx]]
            idx += 1

        chunks.append((chunk, start, idx))

        next_start = idx - overlap
        if next_start <= start:
            next_start = idx
        start = next_start

    return chunks
