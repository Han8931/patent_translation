# translate.py

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import TypedDict, List, Dict, Literal, Optional

from docx import Document
from langgraph.graph import StateGraph, END
from openai import OpenAI

from chunking import (
    Block,
    BlockType,
    iter_blocks,
    chunk_blocks_with_spans,
    build_contexts,
)
from prompts import PROMPTS

# ============================================================
# 0) Types
# ============================================================

Section = Literal["default", "claims", "abstract"]

class TranslateState(TypedDict):
    chunks: List[List[Block]]
    contexts: List[Dict[str, str]]
    i: int
    results: Dict[str, str]

    model: str
    base_url: str | None
    api_key: str | None

    target_lang: str

    # routing / prompts
    section: Section
    prompt_default: str
    prompt_claims: str
    prompt_abstract: str

    # abstract word-count tracking
    abstract_word_count: int
    last_abstract_block_id: str | None
    abstract_completed: bool

    # cross-chunk context for terminology consistency
    glossary: Dict[str, str]       # Korean term → English term (accumulated)
    prev_translated_text: str      # English output of the previous chunk

    # cross-chunk context for claim preambles
    claim_preambles: Dict[str, str]  # claim number → category (e.g. "1" → "method")


# ============================================================
# 1) OpenAI client + JSON extraction fallback
# ============================================================

def make_client(base_url: Optional[str] = None, api_key: Optional[str] = None) -> OpenAI:
    kwargs = {"timeout": 120.0, "max_retries": 2}
    if base_url:
        kwargs["base_url"] = base_url
    if api_key:
        kwargs["api_key"] = api_key
    return OpenAI(**kwargs)


def extract_first_json_object(s: str) -> Optional[str]:
    s2 = (s or "").strip()
    if s2.startswith("{") and s2.endswith("}"):
        return s2

    start = s2.find("{")
    if start == -1:
        return None

    depth = 0
    for i in range(start, len(s2)):
        ch = s2[i]
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return s2[start : i + 1]
    return None


# ============================================================
# 2) Translate one chunk with a given prompt
# ============================================================

def _format_glossary(glossary: Dict[str, str]) -> str:
    if not glossary:
        return "(none yet — this is the first chunk)"
    lines = [f"{ko} → {en}" for ko, en in glossary.items()]
    return "\n".join(lines)


def _format_claim_preambles(preambles: Dict[str, str]) -> str:
    if not preambles:
        return "(none yet — no independent claims translated so far)"
    lines = [f"Claim {num}: {cat}" for num, cat in sorted(preambles.items(), key=lambda x: int(x[0]) if x[0].isdigit() else 0)]
    return "\n".join(lines)


_INDEP_CLAIM_RE = re.compile(
    r"^\s*(\d+)\.\s+An?\s+(\w+)",
    re.MULTILINE,
)

_CATEGORY_WORDS = {"method", "apparatus", "system", "device", "medium", "program", "composition", "compound", "kit", "assembly", "machine", "article", "process", "circuit", "network", "computer", "storage", "structure", "sensor", "module", "unit", "arrangement"}

def extract_claim_preambles(translated_text: str) -> Dict[str, str]:
    """Regex fallback: extract independent claim preambles from translated English text."""
    preambles: Dict[str, str] = {}
    for m in _INDEP_CLAIM_RE.finditer(translated_text):
        claim_num = m.group(1)
        category = m.group(2).lower()
        if category in _CATEGORY_WORDS:
            preambles[claim_num] = category
    return preambles


def translate_chunk(
    client: OpenAI,
    model: str,
    *,
    prompt_name: str,
    target_lang: str,
    chunk: List[Block],
    context_before: str = "",
    context_after: str = "",
    glossary: Dict[str, str] | None = None,
    prev_translated_text: str = "",
    claim_preambles: Dict[str, str] | None = None,
) -> tuple[Dict[str, str], List[Dict[str, str]], Dict[str, str]]:
    """Translate a chunk and return (translations, key_terms, new_claim_preambles)."""
    prompt = PROMPTS[prompt_name]
    payload = [{"id": b.id, "text": b.text} for b in chunk]

    glossary_str = _format_glossary(glossary or {})
    prev_translation = prev_translated_text or "(none — this is the first chunk)"
    preambles_str = _format_claim_preambles(claim_preambles or {})

    system_prompt = prompt.system
    user_prompt = prompt.render_user(
        payload=payload,
        target_lang=target_lang,
        context_before=context_before,
        context_after=context_after,
        glossary=glossary_str,
        prev_translation=prev_translation,
        claim_preambles=preambles_str,
    )

    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.0,
    )

    content = (resp.choices[0].message.content or "").strip()

    try:
        data = json.loads(content)
    except Exception:
        json_str = extract_first_json_object(content)
        if not json_str:
            raise ValueError(f"Model did not return JSON. Got:\n{content[:800]}")
        data = json.loads(json_str)

    out: Dict[str, str] = {}
    translations = data.get("translations", [])
    if isinstance(translations, list):
        for item in translations:
            if isinstance(item, dict) and "id" in item and "text" in item:
                out[str(item["id"])] = str(item["text"])

    # Extract key_terms for glossary accumulation
    key_terms: List[Dict[str, str]] = []
    raw_terms = data.get("key_terms", [])
    if isinstance(raw_terms, list):
        for term in raw_terms:
            if isinstance(term, dict) and "ko" in term and "en" in term:
                key_terms.append({"ko": str(term["ko"]), "en": str(term["en"])})

    # Extract claim_preambles from LLM response (primary)
    new_preambles: Dict[str, str] = {}
    raw_preambles = data.get("claim_preambles", [])
    if isinstance(raw_preambles, list):
        for p in raw_preambles:
            if isinstance(p, dict) and "claim_num" in p and "category" in p:
                new_preambles[str(p["claim_num"])] = str(p["category"]).lower()

    # Regex fallback: extract from translated text if LLM didn't provide them
    if not new_preambles and out:
        translated_text = "\n".join(out.values())
        new_preambles = extract_claim_preambles(translated_text)

    return out, key_terms, new_preambles


# ============================================================
# 3) Abstract word counting + finalization (NEW)
# ============================================================

_WORD_RE = re.compile(r"[A-Za-z]+(?:'[A-Za-z]+)?")

def count_english_words(text: str) -> int:
    return len(_WORD_RE.findall(text or ""))

def append_word_count_to_last_sentence(text: str, n: int) -> str:
    s = (text or "").rstrip()
    if not s:
        return s
    if re.search(r"\(\d+\s+words\)\s*$", s):
        return s
    return f"{s} ({n} words)"


# ============================================================
# 4) Section detection + routing (UPDATED headings)
# ============================================================

_ABSTRACT_HEADING_RE = re.compile(r"【?\s*요\s*약\s*】?|요약서|초록|ABSTRACT", re.IGNORECASE)
_CLAIMS_HEADING_RE = re.compile(r"【?\s*청구\s*범위\s*】?|청구항|CLAIMS?", re.IGNORECASE)


def _chunk_text(chunk: List[Block]) -> str:
    return "\n".join((b.text or "") for b in chunk)


def detect_section_from_chunk(prev: Section, chunk: List[Block]) -> Section:
    """
    Detect section headings with a few common variants:
    - Claims: 【청구 범위】, 【청구범위】, 청구항, Claims/CLAIMS
    - Abstract: 【요약】 plus loose spacing, 요약서, 초록, ABSTRACT
    Sticky behavior: if no heading is found, keep previous section.
    """
    t = _chunk_text(chunk)

    m_abs = _ABSTRACT_HEADING_RE.search(t)
    m_claim = _CLAIMS_HEADING_RE.search(t)

    if m_abs and m_claim:
        # Choose the heading that appears first in the chunk
        return "abstract" if m_abs.start() <= m_claim.start() else "claims"
    if m_abs:
        return "abstract"
    if m_claim:
        return "claims"
    return prev

def node_route_section(state: TranslateState) -> TranslateState:
    i = state["i"]

    def finalize_abstract_word_count(s: TranslateState) -> TranslateState:
        n = int(s.get("abstract_word_count", 0) or 0)
        last_id = s.get("last_abstract_block_id")
        results = s.get("results", {})
        if n > 0 and last_id and last_id in results:
            new_results = dict(results)
            new_results[last_id] = append_word_count_to_last_sentence(new_results[last_id], n)
            return {**s, "results": new_results, "abstract_completed": True}
        return s

    # If we are done, finalize abstract once here
    if i >= len(state["chunks"]):
        if state.get("section") == "abstract" or state.get("abstract_word_count", 0):
            return finalize_abstract_word_count(state)
        return state

    prev = state.get("section", "default")
    new = detect_section_from_chunk(prev, state["chunks"][i])

    # If we've already processed abstract content, force exit after first abstract body chunk
    if prev == "abstract" and state.get("abstract_completed"):
        new = "default"

    # Leaving abstract -> finalize "(XXX words)" on last abstract sentence
    if prev == "abstract" and new != "abstract":
        finalized = finalize_abstract_word_count(state)
        if finalized is not state:
            print(f"[ABSTRACT] finalize word count: {finalized.get('abstract_word_count', 0)} words")
            state = finalized

    if new != prev:
        print(f"[ROUTE] chunk {i}: {prev} -> {new}")

    return {**state, "section": new}


def route_after_router(state: TranslateState) -> str:
    if state["i"] >= len(state["chunks"]):
        return END

    sec = state.get("section", "default")
    if sec == "claims":
        return "classify_claims"
    if sec == "abstract":
        return "translate_abstract"
    return "translate_default"


# ============================================================
# 5) Translation nodes (default/claims/abstract)
# ============================================================

def _translate_with_prompt(state: TranslateState, prompt_name: str) -> TranslateState:
    i = state["i"]
    chunks = state["chunks"]
    contexts = state.get("contexts", [])
    if i >= len(chunks):
        return state

    client = make_client(state.get("base_url"), state.get("api_key"))
    chunk = chunks[i]
    ctx_before = contexts[i].get("before", "") if i < len(contexts) else ""
    ctx_after = contexts[i].get("after", "") if i < len(contexts) else ""
    sec = state.get("section", "default")
    glossary = dict(state.get("glossary", {}))
    prev_translated_text = state.get("prev_translated_text", "")
    claim_preambles = dict(state.get("claim_preambles", {}))

    print(f"[{sec}] {i}-th chunk... (glossary: {len(glossary)} terms, preambles: {len(claim_preambles)})")

    translated_map, key_terms, new_preambles = translate_chunk(
        client=client,
        model=state["model"],
        prompt_name=prompt_name,
        target_lang=state["target_lang"],
        chunk=chunk,
        context_before=ctx_before,
        context_after=ctx_after,
        glossary=glossary,
        prev_translated_text=prev_translated_text,
        claim_preambles=claim_preambles,
    )

    results = dict(state["results"])
    results.update(translated_map)

    # Accumulate glossary with newly extracted terms
    for term in key_terms:
        glossary[term["ko"]] = term["en"]

    # Accumulate claim preambles
    claim_preambles.update(new_preambles)

    # Build prev_translated_text from this chunk's English output
    new_prev = "\n".join(
        translated_map.get(b.id, b.text) for b in chunk
    )

    # Track abstract word count
    abstract_word_count = int(state.get("abstract_word_count", 0))
    last_abstract_block_id = state.get("last_abstract_block_id")
    abstract_completed = state.get("abstract_completed", False)

    if sec == "abstract":
        for bid, txt in translated_map.items():
            t = (txt or "").strip()
            # Drop leading heading if the model left it attached.
            t_no_heading = re.sub(r"^ABSTRACT[:\s\-]*", "", t, flags=re.IGNORECASE)
            if t_no_heading.upper() == "ABSTRACT":
                continue
            wc = count_english_words(t_no_heading)
            if wc > 0:
                abstract_word_count += wc
                last_abstract_block_id = bid
        if abstract_word_count > 0:
            abstract_completed = True

    return {
        **state,
        "results": results,
        "i": i + 1,
        "abstract_word_count": abstract_word_count,
        "last_abstract_block_id": last_abstract_block_id,
        "abstract_completed": abstract_completed,
        "glossary": glossary,
        "prev_translated_text": new_prev,
        "claim_preambles": claim_preambles,
    }


def node_translate_default(state: TranslateState) -> TranslateState:
    return _translate_with_prompt(state, state["prompt_default"])

def node_translate_claims(state: TranslateState) -> TranslateState:
    return _translate_with_prompt(state, state["prompt_claims"])


# New: classify claim chunks as independent/dependent and route to specialized prompts
def classify_claims(state: TranslateState) -> str:
    i = state.get("i", 0)
    if i >= len(state.get("chunks", [])):
        return "translate_claims_indep"  # fallback

    chunk = state["chunks"][i]
    text = "\n".join((b.text or "") for b in chunk)

    # Simple heuristic: if the chunk references another claim (제 <num> 항 or "claim <num>") treat as dependent
    if re.search(r"제\s*\d+\s*항", text) or re.search(r"claim\s+\d+", text, re.IGNORECASE):
        return "translate_claims_dep"
    return "translate_claims_indep"


def node_translate_claims_indep(state: TranslateState) -> TranslateState:
    return _translate_with_prompt(state, state.get("prompt_claims_indep", "patent_kr2en_claims_indep_v1"))


def node_translate_claims_dep(state: TranslateState) -> TranslateState:
    return _translate_with_prompt(state, state.get("prompt_claims_dep", "patent_kr2en_claims_dep_v1"))

def node_translate_abstract(state: TranslateState) -> TranslateState:
    return _translate_with_prompt(state, state["prompt_abstract"])


# ============================================================
# 6) Build graph (router -> specialized translate -> router ...)
# ============================================================

def build_translation_graph():
    g = StateGraph(TranslateState)

    g.add_node("route_section", node_route_section)
    g.add_node("translate_default", node_translate_default)
    g.add_node("classify_claims", classify_claims)
    g.add_node("translate_claims", node_translate_claims)  # legacy fallback
    g.add_node("translate_claims_indep", node_translate_claims_indep)
    g.add_node("translate_claims_dep", node_translate_claims_dep)
    g.add_node("translate_abstract", node_translate_abstract)

    g.set_entry_point("route_section")

    g.add_conditional_edges(
        "route_section",
        route_after_router,
        {
            "translate_default": "translate_default",
            "classify_claims": "classify_claims",
            "translate_claims": "translate_claims",
            "translate_claims_indep": "translate_claims_indep",
            "translate_claims_dep": "translate_claims_dep",
            "translate_abstract": "translate_abstract",
            END: END,
        },
    )

    g.add_edge("translate_default", "route_section")
    g.add_edge("translate_claims", "route_section")
    g.add_edge("translate_claims_indep", "route_section")
    g.add_edge("translate_claims_dep", "route_section")
    g.add_edge("translate_abstract", "route_section")

    return g.compile()


# ============================================================
# 7) Apply translations back into DOCX
# ============================================================

def apply_translations_to_docx(
    src_docx_path: str | Path,
    translations: Dict[str, str],
    out_docx_path: str | Path,
) -> None:
    doc = Document(str(src_docx_path))

    for block_id, translated in translations.items():
        if block_id.startswith("p:"):
            _, idx = block_id.split(":", 1)
            pi = int(idx)
            if 0 <= pi < len(doc.paragraphs):
                doc.paragraphs[pi].text = translated

        elif block_id.startswith("cell:"):
            parts = block_id.split(":")
            if len(parts) == 5:
                _, ti, ri, ci, pi = parts
                ti = int(ti); ri = int(ri); ci = int(ci); pi = int(pi)

                if 0 <= ti < len(doc.tables):
                    table = doc.tables[ti]
                    if 0 <= ri < len(table.rows):
                        row = table.rows[ri]
                        if 0 <= ci < len(row.cells):
                            cell = row.cells[ci]
                            if 0 <= pi < len(cell.paragraphs):
                                cell.paragraphs[pi].text = translated

    doc.save(str(out_docx_path))


# ============================================================
# 8) End-to-end runner
# ============================================================

def translate_docx(
    src_docx_path: str | Path,
    out_docx_path: str | Path,
    *,
    model: str,
    target_lang: str = "English",
    max_chars_per_chunk: int = 2500,
    context_chars: int = 1000,
    chunk_overlap: int = 0,
    base_url: Optional[str] = None,
    api_key: Optional[str] = None,
    prompt_default: str = "patent_kr2en_body_v1",
    prompt_claims: str = "patent_kr2en_claims_v1",
    prompt_claims_indep: str = "patent_kr2en_claims_indep_v1",
    prompt_claims_dep: str = "patent_kr2en_claims_dep_v1",
    prompt_abstract: str = "patent_kr2en_abstract_v1",
) -> None:
    print("Start chunking...")
    blocks = iter_blocks(src_docx_path)
    chunk_spans = chunk_blocks_with_spans(
        blocks,
        max_chars=max_chars_per_chunk,
        overlap=chunk_overlap,
    )
    chunks = [c for (c, _, _) in chunk_spans]
    contexts_raw = build_contexts(
        blocks,
        [(s, e) for (_, s, e) in chunk_spans],
        context_chars=context_chars,
    )
    contexts = [{"before": b, "after": a} for (b, a) in contexts_raw]
    print(f"Finish chunking: {len(chunks)} chunks (context ±{context_chars} chars)")

    print("Build graph...")
    app = build_translation_graph()

    print("Start invoking...")
    final = app.invoke(
        {
            "chunks": chunks,
            "contexts": contexts,
            "i": 0,
            "results": {},
            "model": model,
            "base_url": base_url,
            "api_key": api_key,
            "target_lang": target_lang,

            "section": "default",
            "prompt_default": prompt_default,
            "prompt_claims": prompt_claims,
            "prompt_claims_indep": prompt_claims_indep,
            "prompt_claims_dep": prompt_claims_dep,
            "prompt_abstract": prompt_abstract,

            "abstract_word_count": 0,
            "last_abstract_block_id": None,
            "abstract_completed": False,

            # cross-chunk context for terminology consistency
            "glossary": {},
            "prev_translated_text": "",
            "claim_preambles": {},
        },
        config={"recursion_limit": max(50, len(chunks) * 3)},
    )

    # If the document ends while still in abstract (or word count exists), finalize again defensively
    if final.get("abstract_word_count", 0) or final.get("section") == "abstract":
        final = node_route_section({**final, "i": len(chunks)})


    print("Apply translations...")
    apply_translations_to_docx(src_docx_path, final["results"], out_docx_path)
    print("Done.")
