# translate.py

from __future__ import annotations

import json
import random
import re
import time
from pathlib import Path
from datetime import datetime
from typing import TypedDict, List, Dict, Literal, Optional, Tuple
from json import JSONDecodeError

from docx import Document
from langgraph.graph import StateGraph, END
from openai import OpenAI

from app.batch_tools.chunking import (
    Block,
    BlockType,
    iter_blocks,
    chunk_blocks_with_spans,
    chunk_blocks_patent_with_spans,
    build_contexts,
)
from app.batch_tools.prompts import PROMPTS
from app.batch_tools.post_process import _format_claim_linebreaks

# ============================================================
# 0) Types
# ============================================================

# NEW: add explicit "spec" (specification/description)
Section = Literal["default", "spec", "claims", "abstract"]


class TranslateState(TypedDict):
    chunks: List[List[Block]]
    contexts: List[Dict[str, str]]
    i: int
    results: Dict[str, str]

    model: str
    base_url: str | None
    api_key: str | None

    target_lang: str

    # routing / prompts
    section: Section
    prompt_default: str
    prompt_claims: str
    prompt_claims_indep: str
    prompt_claims_dep: str
    prompt_abstract: str

    # abstract word-count tracking
    abstract_word_count: int
    last_abstract_block_id: str | None
    abstract_completed: bool

    # cross-chunk context for terminology consistency
    glossary: Dict[str, str]  # Korean term → English term (accumulated)
    prev_translated_text: str  # English output of the previous chunk

    # cross-chunk context for claim preambles
    claim_preambles: Dict[str, str]  # claim number → category (e.g. "1" → "method")

    # id-shape quality tracking (mismatch after retry attempts)
    id_mismatch_detected: bool
    id_mismatch_chunks: List[str]
    # quality tracking: empty model output replaced with source text
    empty_fallback_count: int
    empty_fallback_chunks: List[str]


def node_postprocess_claims(state: TranslateState) -> TranslateState:
    """
    Post-process the most recently translated claims chunk:
    - enforce newline after ':' and ';'
    - (optional) indent lines after the first line
    """
    # Only run for claims
    if state.get("section") != "claims":
        return state

    i = int(state.get("i", 0))
    prev_i = i - 1
    chunks = state.get("chunks", [])
    if prev_i < 0 or prev_i >= len(chunks):
        return state

    chunk = chunks[prev_i]
    results = state.get("results", {})
    if not chunk or not results:
        return state

    new_results = dict(results)

    # Apply formatting to any non-empty text in this chunk
    for b in chunk:
        bid = b.id
        txt = (new_results.get(bid) or "")
        if txt.strip():
            new_results[bid] = _format_claim_linebreaks(txt, indent="    ")  # indent optional

    return {**state, "results": new_results}


# ============================================================
# 1) OpenAI client + JSON extraction fallback
# ============================================================


def make_client(base_url: Optional[str] = None, api_key: Optional[str] = None) -> OpenAI:
    kwargs = {"timeout": 2400, "max_retries": 2}
    if base_url:
        kwargs["base_url"] = base_url
    if api_key:
        kwargs["api_key"] = api_key
    return OpenAI(**kwargs)


def repair_json_with_llm(client: OpenAI, model: str, bad_output: str) -> str:
    """
    Ask the model to convert malformed JSON-ish output into valid JSON.
    Keeps content as-is as much as possible; only fixes JSON syntax.
    """
    sys = (
        "You are a strict JSON repair tool.\n"
        "Convert the user's content into VALID JSON that matches the required schema.\n"
        "Rules:\n"
        "- Output ONLY JSON. No markdown. No commentary.\n"
        "- Use DOUBLE QUOTES for all keys/strings.\n"
        "- Do NOT add or remove translation items unless necessary to make JSON valid.\n"
        "- Preserve ids and translated text as-is whenever possible.\n"
        "- If there are unescaped quotes inside text fields, escape them.\n"
        "- Return a single JSON object.\n"
    )
    user = (
        "Fix the following into valid JSON.\n"
        'Required top-level keys: "translations" (list of {"id","text"})\n'
        'Optional keys: "key_terms", "claim_preambles"\n'
        "\n"
        "CONTENT TO REPAIR:\n"
        f"{bad_output}\n"
    )

    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": sys},
            {"role": "user", "content": user},
        ],
        temperature=0.0,
    )
    return (resp.choices[0].message.content or "").strip()


def extract_first_json_object(s: str) -> Optional[str]:
    s2 = (s or "").strip()
    if s2.startswith("{") and s2.endswith("}"):
        return s2

    start = s2.find("{")
    if start == -1:
        return None

    depth = 0
    for i in range(start, len(s2)):
        ch = s2[i]
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return s2[start : i + 1]
    return None


# ============================================================
# 1.5) Claims post-processing helpers (NEW)
# ============================================================

# Detect claim number from Korean headers like 【청구항 19】 / 청구항 19
_CLAIM_NUM_FROM_BRACKET = re.compile(r"【\s*청구항\s*(\d+)\s*】")
_CLAIM_NUM_FROM_KR = re.compile(r"\b청구항\s*(\d+)\b")
# English numeric prefix: "19. "
_CLAIM_EN_PREFIX = re.compile(r"^\s*(\d+)\.\s+")
# English "CLAIM 19" / "Claim 19"
_CLAIM_EN_WORDING = re.compile(r"^\s*CLAIM\s+(\d+)\b[:\s\-]*", re.IGNORECASE)


def _extract_claim_number_from_chunk(chunk: List[Block]) -> Optional[str]:
    for b in chunk:
        t = (b.text or "").strip()
        m = _CLAIM_NUM_FROM_BRACKET.search(t)
        if m:
            return m.group(1)
        m = _CLAIM_NUM_FROM_KR.search(t)
        if m:
            return m.group(1)
    return None


def _ensure_claim_number_prefix(text: str, claim_num: str) -> str:
    s = (text or "").lstrip()
    if not claim_num:
        return text or ""

    # Already "N. ..."
    if _CLAIM_EN_PREFIX.match(s):
        return s

    # Strip "CLAIM N" / "Claim N" prefixes if present
    s = re.sub(r"^\s*CLAIM\s+%s\b[:\s\-]*" % re.escape(claim_num), "", s, flags=re.IGNORECASE)
    s = re.sub(r"^\s*Claim\s+%s\b[:\s\-]*" % re.escape(claim_num), "", s, flags=re.IGNORECASE)

    return f"{claim_num}. {s}".strip()


def _move_first_nonempty_to_first_id(translated_map: Dict[str, str], chunk: List[Block]) -> Dict[str, str]:
    """
    If the model violated the merge rule and put the merged claim text into a later item,
    move the first non-empty string to the first block id and blank out the source.
    """
    if not chunk:
        return translated_map

    first_id = chunk[0].id
    first_txt = (translated_map.get(first_id) or "").strip()
    if first_txt:
        return translated_map

    # Find first non-empty among the chunk ids
    for b in chunk[1:]:
        txt = (translated_map.get(b.id) or "").strip()
        if txt:
            new_map = dict(translated_map)
            new_map[first_id] = txt
            new_map[b.id] = ""
            return new_map
    return translated_map


# ============================================================
# 2) Translate one chunk with a given prompt
# ============================================================


def _format_glossary(glossary: Dict[str, str]) -> str:
    if not glossary:
        return "(none yet — this is the first chunk)"
    lines = [f"{ko} → {en}" for ko, en in glossary.items()]
    return "\n".join(lines)


def _format_claim_preambles(preambles: Dict[str, str]) -> str:
    if not preambles:
        return "(none yet — no independent claims translated so far)"
    lines = [
        f"Claim {num}: {cat}"
        for num, cat in sorted(
            preambles.items(),
            key=lambda x: int(x[0]) if x[0].isdigit() else 0,
        )
    ]
    return "\n".join(lines)


_INDEP_CLAIM_RE = re.compile(r"^\s*(\d+)\.\s+An?\s+(\w+)", re.MULTILINE)

_CATEGORY_WORDS = {
    "method",
    "apparatus",
    "system",
    "device",
    "medium",
    "program",
    "composition",
    "compound",
    "kit",
    "assembly",
    "machine",
    "article",
    "process",
    "circuit",
    "network",
    "computer",
    "storage",
    "structure",
    "sensor",
    "module",
    "unit",
    "arrangement",
}


def extract_claim_preambles(translated_text: str) -> Dict[str, str]:
    """Regex fallback: extract independent claim preambles from translated English text."""
    preambles: Dict[str, str] = {}
    for m in _INDEP_CLAIM_RE.finditer(translated_text):
        claim_num = m.group(1)
        category = m.group(2).lower()
        if category in _CATEGORY_WORDS:
            preambles[claim_num] = category
    return preambles


def translate_chunk(
    client: OpenAI,
    model: str,
    *,
    prompt_name: str,
    target_lang: str,
    chunk: List[Block],
    context_before: str = "",
    context_after: str = "",
    glossary: Dict[str, str] | None = None,
    prev_translated_text: str = "",
    claim_preambles: Dict[str, str] | None = None,
    retry_instruction: str = "",
    temperature: float = 0.0,  # Fix A: caller controls temperature per attempt
) -> tuple[Dict[str, str], List[Dict[str, str]], Dict[str, str]]:
    """Translate a chunk and return (translations, key_terms, new_claim_preambles)."""
    prompt = PROMPTS[prompt_name]

    # Option 1: Position-based mapping — use sequential integers (0, 1, 2...) instead of
    # opaque string IDs like "p:12" or "cell:0:1:2:3". The LLM only needs to reproduce
    # small integers, which eliminates almost all ID reproduction errors.
    # We remap back to original string IDs after parsing the response.
    pos_to_id = {i: b.id for i, b in enumerate(chunk)}
    payload = [{"id": i, "text": b.text} for i, b in enumerate(chunk)]

    glossary_str = _format_glossary(glossary or {})
    prev_translation = prev_translated_text or "(none — this is the first chunk)"
    preambles_str = _format_claim_preambles(claim_preambles or {})

    system_prompt = prompt.system
    user_prompt = prompt.render_user(
        payload=payload,
        target_lang=target_lang,
        context_before=context_before,
        context_after=context_after,
        glossary=glossary_str,
        prev_translation=prev_translation,
        claim_preambles=preambles_str,
    )

    # Fix E: Inject the required integer positions at the top of every call.
    n = len(chunk)
    ids_note = (
        f"REQUIRED OUTPUT IDs — return EXACTLY {n} items with integer ids 0 to {n - 1}, no more, no less:\n"
        f"{json.dumps(list(range(n)))}\n\n"
    )

    # Fix C: Retry instruction at the TOP so it gets maximum attention.
    if retry_instruction:
        user_prompt = (
            "⚠️ RETRY — CRITICAL FIXES REQUIRED FROM YOUR PREVIOUS ATTEMPT:\n"
            f"{retry_instruction}\n\n"
            + ids_note
            + user_prompt
        )
    else:
        user_prompt = ids_note + user_prompt

    resp = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=temperature,  # Fix A: use caller-supplied temperature
    )

    content = (resp.choices[0].message.content or "").strip()

    # 1) Try direct JSON
    try:
        data = json.loads(content)
    except JSONDecodeError:
        # 2) Try extracting the first {...}
        json_str = extract_first_json_object(content)
        if json_str:
            try:
                data = json.loads(json_str)
            except JSONDecodeError:
                # 3) Ask model to repair JSON
                repaired = repair_json_with_llm(client, model, content)
                repaired_json = extract_first_json_object(repaired) or repaired
                data = json.loads(repaired_json)
        else:
            # 3) Ask model to repair JSON (even if no braces found)
            repaired = repair_json_with_llm(client, model, content)
            repaired_json = extract_first_json_object(repaired) or repaired
            data = json.loads(repaired_json)

    # Remap integer positions back to original string block IDs.
    # Any position outside 0..N-1 is silently dropped (no "unexpected ID" noise).
    out: Dict[str, str] = {}
    translations = data.get("translations", [])
    if isinstance(translations, list):
        for item in translations:
            if isinstance(item, dict) and "id" in item and "text" in item:
                try:
                    pos = int(item["id"])
                    if pos in pos_to_id:
                        out[pos_to_id[pos]] = str(item["text"])
                except (ValueError, TypeError):
                    pass

    # Extract key_terms for glossary accumulation
    key_terms: List[Dict[str, str]] = []
    raw_terms = data.get("key_terms", [])
    if isinstance(raw_terms, list):
        for term in raw_terms:
            if isinstance(term, dict) and "ko" in term and "en" in term:
                key_terms.append({"ko": str(term["ko"]), "en": str(term["en"])})

    # Extract claim_preambles from LLM response (primary)
    new_preambles: Dict[str, str] = {}
    raw_preambles = data.get("claim_preambles", [])
    if isinstance(raw_preambles, list):
        for p in raw_preambles:
            if isinstance(p, dict) and "claim_num" in p and "category" in p:
                new_preambles[str(p["claim_num"])] = str(p["category"]).lower()

    # Regex fallback: extract from translated text if LLM didn't provide them
    if not new_preambles and out:
        translated_text = "\n".join(out.values())
        new_preambles = extract_claim_preambles(translated_text)

    return out, key_terms, new_preambles


# ============================================================
# 3) Abstract word counting + finalization
# ============================================================

_WORD_RE = re.compile(r"[A-Za-z]+(?:'[A-Za-z]+)?")


def count_english_words(text: str) -> int:
    return len(_WORD_RE.findall(text or ""))


def append_word_count_to_last_sentence(text: str, n: int) -> str:
    s = (text or "").rstrip()
    if not s:
        return s
    if re.search(r"\(\d+\s+words\)\s*$", s):
        return s
    return f"{s} ({n} words)"


# ============================================================
# 4) Section detection + routing
# ============================================================

_ABSTRACT_HEADING_RE = re.compile(r"【?\s*요\s*약\s*】?|요약서|초록|ABSTRACT", re.IGNORECASE)
_CLAIMS_HEADING_RE = re.compile(
    r"(^\s*【?\s*청구\s*범위\s*】?\s*$)|"
    r"(^\s*【?\s*청구항\s*\d+\s*(?:】|[.)])?\s*$)|"
    r"(^\s*CLAIMS?\s*$)",
    re.IGNORECASE | re.MULTILINE,
)

# Spec/description detection
_SPEC_HEADING_RE = re.compile(
    r"(발명의\s*상세한\s*설명|상세한\s*설명|명세서|DESCRIPTION|DETAILED\s+DESCRIPTION)",
    re.IGNORECASE,
)

# Representative figure/drawing — appears after abstract in Korean patents.
# Must be treated as a section break so it doesn't inherit "abstract" and pollute word count.
_REPFIG_HEADING_RE = re.compile(
    r"【?\s*대표\s*도\s*】?|대표\s*도면|REPRESENTATIVE\s+(?:DRAWING|FIGURE)",
    re.IGNORECASE,
)


def _chunk_text(chunk: List[Block]) -> str:
    return "\n".join((b.text or "") for b in chunk)


def _has_abstract_heading(chunk: List[Block]) -> bool:
    return bool(_ABSTRACT_HEADING_RE.search(_chunk_text(chunk)))


def detect_section_from_chunk(prev: Section, chunk: List[Block]) -> Section:
    """
    Detect section headings with a few common variants:
    - Claims: 【청구 범위】, 【청구범위】, 청구항, Claims/CLAIMS
    - Abstract: 【요약】 plus loose spacing, 요약서, 초록, ABSTRACT
    - Spec: 발명의 상세한 설명, 상세한 설명, 명세서, DESCRIPTION, DETAILED DESCRIPTION

    Sticky behavior: if no heading is found, keep previous section.
    If multiple headings appear in same chunk, choose the earliest one.
    """
    t = _chunk_text(chunk)

    m_abs = _ABSTRACT_HEADING_RE.search(t)
    m_claim = _CLAIMS_HEADING_RE.search(t)
    m_spec = _SPEC_HEADING_RE.search(t)
    m_repfig = _REPFIG_HEADING_RE.search(t)

    candidates: List[Tuple[Section, int]] = []
    if m_abs:
        candidates.append(("abstract", m_abs.start()))
    if m_claim:
        candidates.append(("claims", m_claim.start()))
    if m_spec:
        candidates.append(("spec", m_spec.start()))
    if m_repfig:
        # Representative figure is not abstract — treat as default to force a section break
        # so the abstract word count finalizes before this block is processed.
        candidates.append(("default", m_repfig.start()))

    if candidates:
        candidates.sort(key=lambda x: x[1])
        return candidates[0][0]

    return prev


def node_route_section(state: TranslateState) -> TranslateState:
    i = state["i"]

    def finalize_abstract_word_count(s: TranslateState) -> TranslateState:
        n = int(s.get("abstract_word_count", 0) or 0)
        last_id = s.get("last_abstract_block_id")
        results = s.get("results", {})
        if n > 0 and last_id and last_id in results:
            new_results = dict(results)
            new_results[last_id] = append_word_count_to_last_sentence(new_results[last_id], n)
            return {**s, "results": new_results, "abstract_completed": True}
        return s

    if i >= len(state["chunks"]):
        if state.get("section") == "abstract" or state.get("abstract_word_count", 0):
            return finalize_abstract_word_count(state)
        return state

    prev = state.get("section", "default")
    curr_chunk = state["chunks"][i]
    new = detect_section_from_chunk(prev, curr_chunk)

    if prev == "abstract" and new != "abstract":
        finalized = finalize_abstract_word_count(state)
        if finalized is not state:
            print(f"[ABSTRACT] finalize word count: {finalized.get('abstract_word_count', 0)} words")
            state = finalized

    if new != prev:
        print(f"[ROUTE] chunk {i}: {prev} -> {new}")

    return {**state, "section": new}


def route_after_router(state: TranslateState) -> str:
    if state["i"] >= len(state["chunks"]):
        return END

    sec = state.get("section", "default")
    if sec == "claims":
        return "classify_claims"
    if sec == "abstract":
        return "translate_abstract"
    # spec/default -> translate_default (spec uses prompt_default)
    return "translate_default"


# ============================================================
# 4.5) NEW: Section-sized chunking + adaptive fallback
# ============================================================


def chunk_blocks_by_section_with_spans(blocks: List[Block]) -> List[tuple[List[Block], int, int]]:
    """
    Returns [(chunk_blocks, start_idx, end_idx_exclusive), ...]
    where each chunk is an entire contiguous section (abstract/claims/spec/default).
    """
    spans: List[tuple[List[Block], int, int]] = []
    if not blocks:
        return spans

    curr_sec: Section = detect_section_from_chunk("default", [blocks[0]])
    start = 0

    for i in range(1, len(blocks)):
        sec_i = detect_section_from_chunk(curr_sec, [blocks[i]])
        if sec_i != curr_sec:
            spans.append((blocks[start:i], start, i))
            start = i
            curr_sec = sec_i

    spans.append((blocks[start:len(blocks)], start, len(blocks)))
    return spans


def total_chars(chunk: List[Block]) -> int:
    return sum(len(b.text or "") for b in chunk)


def build_section_chunks_with_fallback(
    blocks_list: List[Block],
    *,
    max_section_chars: int,
    max_chars_per_chunk: int,
    chunk_overlap: int,
) -> List[tuple[List[Block], int, int]]:
    """
    1) Try section-sized chunks (whole Abstract / whole Claims / whole Spec).
    2) If any section chunk is too large, fallback to patent-aware chunking
       within that section, preserving global indices.
    """
    section_spans = chunk_blocks_by_section_with_spans(blocks_list)

    final_spans: List[tuple[List[Block], int, int]] = []
    for chunk, s, e in section_spans:
        if total_chars(chunk) <= max_section_chars:
            final_spans.append((chunk, s, e))
            continue

        # fallback: split only inside this oversized section
        print(f"[CHUNK] section too large ({total_chars(chunk)} chars). Fallback split: blocks[{s}:{e}]")
        sub_spans = chunk_blocks_patent_with_spans(
            blocks_list[s:e],
            max_chars=max_chars_per_chunk,
            overlap=chunk_overlap,
            claim_overlap=0,  # recommended: don't overlap across claim sub-splits
        )
        for sub_chunk, sub_s, sub_e in sub_spans:
            final_spans.append((sub_chunk, s + sub_s, s + sub_e))

    return final_spans


# ============================================================
# 5) Translation nodes (default/claims/abstract)
# ============================================================


def _normalize_translations_for_chunk(
    *,
    chunk: List[Block],
    translated_map: Dict[str, str],
    section: Section,
) -> tuple[Dict[str, str], List[str], List[str], List[str]]:
    """
    Ensure exactly one output entry per input block id.
    - Drops unexpected ids from model output.
    - Fills missing ids deterministically.
    """
    expected_ids = [b.id for b in chunk]
    expected_set = set(expected_ids)
    source_by_id = {b.id: (b.text or "") for b in chunk}

    unexpected = [k for k in translated_map.keys() if k not in expected_set]
    missing = [bid for bid in expected_ids if bid not in translated_map]

    normalized: Dict[str, str] = {}
    fallback_ids: List[str] = []
    for bid in expected_ids:
        src_text = source_by_id[bid]
        src_nonempty = bool(src_text.strip())

        if bid in translated_map:
            out_text = str(translated_map[bid] or "")

            # Guardrail: non-empty source should not become empty output.
            # Exception: claims mode intentionally allows empty non-first ids
            # when a merged claim body is stored in the first id.
            if not out_text.strip() and src_nonempty:
                if section == "claims" and bid != expected_ids[0]:
                    normalized[bid] = ""
                else:
                    normalized[bid] = src_text
                    fallback_ids.append(bid)
            else:
                normalized[bid] = out_text
            continue

        # Claims mode intentionally allows empty non-first items due merge rules.
        if section == "claims" and bid != expected_ids[0]:
            normalized[bid] = ""
        else:
            # Prefer preserving source text over silent drops.
            normalized[bid] = source_by_id[bid]
            if src_nonempty:
                fallback_ids.append(bid)

    # Claims guard: first item should never be blank when possible.
    if section == "claims" and expected_ids:
        first_id = expected_ids[0]
        if not (normalized.get(first_id) or "").strip():
            for bid in expected_ids[1:]:
                txt = normalized.get(bid, "")
                if txt.strip():
                    normalized[first_id] = txt
                    normalized[bid] = ""
                    break

            if not (normalized.get(first_id) or "").strip():
                normalized[first_id] = source_by_id[first_id]

    return normalized, missing, unexpected, fallback_ids


def _repair_missing_blocks(
    client: OpenAI,
    model: str,
    chunk: List[Block],
    missing_ids: List[str],
    target_lang: str,
    section: Section,
    glossary: Dict[str, str] | None = None,
    claim_preambles: Dict[str, str] | None = None,
) -> Dict[str, str]:
    """
    Fix D: Targeted repair call — send ONLY the missing blocks with a minimal,
    focused prompt. Much smaller payload → near-100% ID coverage.
    Called after all main retry attempts still have missing IDs.
    """
    missing_set = set(missing_ids)
    repair_blocks = [b for b in chunk if b.id in missing_set]
    if not repair_blocks:
        return {}

    # Position-based mapping for repair call too: 0, 1, 2... → original string IDs
    pos_to_id = {i: b.id for i, b in enumerate(repair_blocks)}
    payload = [{"id": i, "text": b.text} for i, b in enumerate(repair_blocks)]
    n = len(repair_blocks)
    glossary_str = _format_glossary(glossary or {})
    preambles_str = _format_claim_preambles(claim_preambles or {})

    claims_note = ""
    if section == "claims":
        claims_note = (
            "\nSection is CLAIMS. For multi-line claims, put the full rendered claim text "
            "in the first item of each claim and set remaining same-claim items to empty string.\n"
            f"CLAIM PREAMBLES:\n{preambles_str}\n"
        )

    system = (
        "You are a professional patent translation engine (Korean → English).\n"
        "Translate each JSON item faithfully. Output ONLY valid JSON with double quotes.\n"
        "Return exactly one translations item per input id — no omissions, no additions.\n"
    )
    user = (
        f"REQUIRED OUTPUT IDs — return EXACTLY {n} items with integer ids 0 to {n - 1}:\n"
        f"{json.dumps(list(range(n)))}\n\n"
        f"GLOSSARY:\n{glossary_str}\n"
        f"{claims_note}"
        f'Output schema: {{"translations":[{{"id":0,"text":"..."}},{{"id":1,"text":"..."}},...]}}\n\n'
        f"INPUT:\n{json.dumps(payload, ensure_ascii=False)}\n"
    )

    try:
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user},
            ],
            temperature=0.0,
        )
        content = (resp.choices[0].message.content or "").strip()
        try:
            data = json.loads(content)
        except JSONDecodeError:
            json_str = extract_first_json_object(content)
            if not json_str:
                return {}
            data = json.loads(json_str)

        out: Dict[str, str] = {}
        for item in data.get("translations", []):
            if isinstance(item, dict) and "id" in item and "text" in item:
                try:
                    pos = int(item["id"])
                    if pos in pos_to_id:
                        out[pos_to_id[pos]] = str(item["text"])
                except (ValueError, TypeError):
                    pass
        return out
    except Exception as exc:
        print(f"[REPAIR] call failed: {type(exc).__name__}: {exc}")
        return {}


def _translate_with_prompt(state: TranslateState, prompt_name: str) -> TranslateState:
    i = state["i"]
    chunks = state["chunks"]
    contexts = state.get("contexts", [])
    if i >= len(chunks):
        return state

    client = make_client(state.get("base_url"), state.get("api_key"))
    chunk = chunks[i]
    ctx_before = contexts[i].get("before", "") if i < len(contexts) else ""
    ctx_after = contexts[i].get("after", "") if i < len(contexts) else ""
    sec = state.get("section", "default")

    glossary = dict(state.get("glossary", {}))
    prev_translated_text = state.get("prev_translated_text", "")
    claim_preambles = dict(state.get("claim_preambles", {}))
    id_mismatch_detected = bool(state.get("id_mismatch_detected", False))
    id_mismatch_chunks = list(state.get("id_mismatch_chunks", []))
    empty_fallback_count = int(state.get("empty_fallback_count", 0))
    empty_fallback_chunks = list(state.get("empty_fallback_chunks", []))

    # Skip pure-empty chunks to avoid pointless LLM calls and empty JSON responses.
    if all(not (b.text or "").strip() for b in chunk):
        results = dict(state["results"])
        for b in chunk:
            results[b.id] = b.text or ""
        print(f"[SKIP] chunk {i}: all blocks are empty/whitespace")
        return {
            **state,
            "results": results,
            "i": i + 1,
        }

    print(f"[{sec}] {i}-th chunk... (glossary: {len(glossary)} terms, preambles: {len(claim_preambles)})")

    max_attempts = 3
    translated_map: Dict[str, str] = {}
    key_terms: List[Dict[str, str]] = []
    new_preambles: Dict[str, str] = {}
    missing_ids: List[str] = []
    unexpected_ids: List[str] = []
    fallback_ids: List[str] = []
    best_score: Tuple[int, int] = (10**9, 10**9)

    # Fix B: track each attempt's specific missing/unexpected IDs for the next retry instruction.
    # We store them as integer positions (matching what the LLM sees) so the retry instruction
    # is immediately actionable — the model already knows positions, not opaque string IDs.
    id_to_pos = {b.id: i for i, b in enumerate(chunk)}
    prev_attempt_missing_pos: List[int] = []
    prev_attempt_unexpected_pos: List[int] = []

    for attempt in range(1, max_attempts + 1):
        retry_instruction = ""
        if attempt > 1:
            # Fix A: use non-zero temperature on retries so the model can produce different output
            attempt_temperature = 0.3
            # Fix B: include the specific integer positions that were wrong in the previous attempt
            if sec == "claims":
                base_msg = (
                    "Return exactly one translations item per input id. "
                    "Do not add new ids, do not remove ids, and keep each input id exactly once. "
                    "The FIRST item of each claim MUST have non-empty text with the full rendered claim. "
                    "Only set text to empty string for non-first items of a merged multi-item claim."
                )
            else:
                base_msg = (
                    "Return exactly one translations item per input id. "
                    "Do not add new ids, do not remove ids, and keep each input id exactly once. "
                    "For EVERY non-empty input text, you MUST return a non-empty English translation. "
                    "Only return empty text when the input text is actually empty or whitespace-only."
                )
            issues: List[str] = []
            if prev_attempt_missing_pos:
                issues.append(f"Integer ids missing from your last response (you must add these): {prev_attempt_missing_pos}")
            if prev_attempt_unexpected_pos:
                issues.append(f"Integer ids you returned that are out of range (remove these): {prev_attempt_unexpected_pos}")
            retry_instruction = base_msg
            if issues:
                retry_instruction += "\n" + "\n".join(issues)
        else:
            attempt_temperature = 0.0  # Fix A: first attempt stays deterministic

        try:
            cand_map, cand_terms, cand_preambles = translate_chunk(
                client=client,
                model=state["model"],
                prompt_name=prompt_name,
                target_lang=state["target_lang"],
                chunk=chunk,
                context_before=ctx_before,
                context_after=ctx_after,
                glossary=glossary,
                prev_translated_text=prev_translated_text,
                claim_preambles=claim_preambles,
                retry_instruction=retry_instruction,
                temperature=attempt_temperature,  # Fix A
            )
        except Exception as exc:
            print(
                f"[WARN] chunk {i}, attempt {attempt}/{max_attempts}: "
                f"LLM call raised {type(exc).__name__}: {exc}"
            )
            if attempt < max_attempts:
                delay = (0.8 * (2 ** (attempt - 1))) + random.uniform(0.0, 0.25)
                print(f"[RETRY] chunk {i}: retrying after exception in {delay:.2f}s")
                time.sleep(delay)
                continue
            # All attempts exhausted — fall back to source text so the document isn't lost.
            print(
                f"[FAIL] chunk {i}: all {max_attempts} attempts raised exceptions. "
                f"Falling back to source text for this chunk."
            )
            cand_map = {b.id: (b.text or "") for b in chunk}
            cand_terms = []
            cand_preambles = {}

        # -------------------------
        # Claims: enforce merge + numbering (NEW)
        # -------------------------
        if sec == "claims":
            cand_map = _move_first_nonempty_to_first_id(cand_map, chunk)

            claim_num = _extract_claim_number_from_chunk(chunk)
            if claim_num and chunk:
                first_id = chunk[0].id
                cand_map[first_id] = _ensure_claim_number_prefix(cand_map.get(first_id, ""), claim_num)

        cand_map, cand_missing, cand_unexpected, cand_fallback_ids = _normalize_translations_for_chunk(
            chunk=chunk,
            translated_map=cand_map,
            section=sec,
        )
        if sec == "claims":
            # Only the first block of a claim matters — non-first empties are intentional (merge rule).
            first_id = chunk[0].id if chunk else None
            critical_fallbacks = 1 if (first_id and first_id in cand_fallback_ids) else 0
        else:
            critical_fallbacks = len(cand_fallback_ids)
        cand_score = (
            len(cand_missing) + len(cand_unexpected) + critical_fallbacks,
            len(cand_unexpected),
        )
        if cand_score < best_score:
            best_score = cand_score
            translated_map = cand_map
            key_terms = cand_terms
            new_preambles = cand_preambles
            missing_ids = cand_missing
            unexpected_ids = cand_unexpected
            fallback_ids = cand_fallback_ids

        if not cand_missing and not cand_unexpected and critical_fallbacks == 0:
            break

        # Fix B: convert missing/unexpected string IDs → integer positions for the next retry instruction
        prev_attempt_missing_pos = [id_to_pos[bid] for bid in cand_missing if bid in id_to_pos]
        # With position-based mapping, unexpected IDs from the model are already filtered at parse time.
        # But if any slip through normalization, convert them too (best-effort).
        prev_attempt_unexpected_pos = [id_to_pos[bid] for bid in cand_unexpected if bid in id_to_pos]

        if attempt < max_attempts:
            delay = (0.8 * (2 ** (attempt - 1))) + random.uniform(0.0, 0.25)
            print(
                f"[RETRY] chunk {i}: id mismatch on attempt {attempt}/{max_attempts} "
                f"(missing={len(cand_missing)}, unexpected={len(cand_unexpected)}, "
                f"empty_fallback={critical_fallbacks}). "
                f"retrying in {delay:.2f}s"
            )
            time.sleep(delay)

    # Fix D: targeted repair call for blocks that are still missing after all main attempts.
    # Only re-sends the missing blocks with a small focused prompt → high ID-coverage success rate.
    if missing_ids:
        print(f"[REPAIR] chunk {i}: targeted repair for {len(missing_ids)} still-missing block(s): {missing_ids}")
        repaired = _repair_missing_blocks(
            client=client,
            model=state["model"],
            chunk=chunk,
            missing_ids=missing_ids,
            target_lang=state["target_lang"],
            section=sec,
            glossary=glossary,
            claim_preambles=claim_preambles,
        )
        recovered = []
        for bid, txt in repaired.items():
            if txt.strip():
                translated_map[bid] = txt
                recovered.append(bid)
        if recovered:
            missing_ids = [bid for bid in missing_ids if bid not in recovered]
            print(f"[REPAIR] chunk {i}: recovered {len(recovered)} block(s): {recovered}")

    if missing_ids or unexpected_ids:
        print(
            f"[WARN] chunk {i}: normalized translation ids "
            f"(missing={len(missing_ids)}, unexpected={len(unexpected_ids)})"
        )
        id_mismatch_detected = True
        id_mismatch_chunks.append(
            f"chunk={i}, section={sec}, missing={len(missing_ids)}, unexpected={len(unexpected_ids)}"
        )

    if fallback_ids:
        print(
            f"[WARN] chunk {i}: empty translation fallback applied to "
            f"{len(fallback_ids)} block(s) in section '{sec}' — source text preserved"
        )
        empty_fallback_count += len(fallback_ids)
        empty_fallback_chunks.append(
            f"chunk={i}, section={sec}, blocks={len(fallback_ids)}"
        )

    results = dict(state["results"])
    results.update(translated_map)

    # Accumulate glossary with newly extracted terms
    for term in key_terms:
        glossary[term["ko"]] = term["en"]

    # Accumulate claim preambles
    claim_preambles.update(new_preambles)

    # Build prev_translated_text from this chunk's English output
    new_prev = "\n".join(translated_map.get(b.id, b.text) for b in chunk)

    # Track abstract word count
    abstract_word_count = int(state.get("abstract_word_count", 0))
    last_abstract_block_id = state.get("last_abstract_block_id")
    abstract_completed = state.get("abstract_completed", False)

    if sec == "abstract":
        for bid, txt in translated_map.items():
            t = (txt or "").strip()
            t_no_heading = re.sub(r"^ABSTRACT[:\s\-]*", "", t, flags=re.IGNORECASE)
            if t_no_heading.upper() == "ABSTRACT":
                continue
            wc = count_english_words(t_no_heading)
            if wc > 0:
                abstract_word_count += wc
                last_abstract_block_id = bid
        if abstract_word_count > 0:
            abstract_completed = True

    return {
        **state,
        "results": results,
        "i": i + 1,
        "abstract_word_count": abstract_word_count,
        "last_abstract_block_id": last_abstract_block_id,
        "abstract_completed": abstract_completed,
        "glossary": glossary,
        "prev_translated_text": new_prev,
        "claim_preambles": claim_preambles,
        "id_mismatch_detected": id_mismatch_detected,
        "id_mismatch_chunks": id_mismatch_chunks,
        "empty_fallback_count": empty_fallback_count,
        "empty_fallback_chunks": empty_fallback_chunks,
    }


def node_translate_default(state: TranslateState) -> TranslateState:
    return _translate_with_prompt(state, state["prompt_default"])


def node_translate_claims(state: TranslateState) -> TranslateState:
    return _translate_with_prompt(state, state["prompt_claims"])


# --- CLAIM CLASSIFICATION ROUTING (FIXED) ---


def node_classify_claims(state: TranslateState) -> TranslateState:
    return state


def route_after_claim_classify(state: TranslateState) -> str:
    i = state.get("i", 0)
    if i >= len(state.get("chunks", [])):
        return "translate_claims_indep"

    chunk = state["chunks"][i]
    text = "\n".join((b.text or "") for b in chunk)

    has_dep = bool(re.search(r"제\s*\d+\s*항", text) or re.search(r"claim\s+\d+", text, re.IGNORECASE))
    has_indep = bool(
        re.search(r"【\s*청구항\s*\d+\s*】", text)
        or re.search(r"\b청구항\s*\d+\b", text)
        or re.search(r"^\s*\d+\s*[.)]", text, re.MULTILINE)
    )

    if has_dep and has_indep:
        return "translate_claims"
    if has_dep:
        return "translate_claims_dep"
    return "translate_claims_indep"


def node_translate_claims_indep(state: TranslateState) -> TranslateState:
    return _translate_with_prompt(
        state,
        state.get("prompt_claims_indep", "patent_kr2en_claims_indep_v1"),
    )


def node_translate_claims_dep(state: TranslateState) -> TranslateState:
    return _translate_with_prompt(
        state,
        state.get("prompt_claims_dep", "patent_kr2en_claims_dep_v1"),
    )


def node_translate_abstract(state: TranslateState) -> TranslateState:
    return _translate_with_prompt(state, state["prompt_abstract"])


# ============================================================
# 6) Build graph (router -> specialized translate -> router ...)
# ============================================================


def build_translation_graph():
    g = StateGraph(TranslateState)

    g.add_node("route_section", node_route_section)
    g.add_node("translate_default", node_translate_default)

    g.add_node("postprocess_claims", node_postprocess_claims)

    g.add_node("classify_claims", node_classify_claims)

    g.add_node("translate_claims", node_translate_claims)  # legacy fallback
    g.add_node("translate_claims_indep", node_translate_claims_indep)
    g.add_node("translate_claims_dep", node_translate_claims_dep)
    g.add_node("translate_abstract", node_translate_abstract)

    g.set_entry_point("route_section")

    g.add_conditional_edges(
        "route_section",
        route_after_router,
        {
            "translate_default": "translate_default",
            "classify_claims": "classify_claims",
            "translate_abstract": "translate_abstract",
            END: END,
        },
    )

    g.add_conditional_edges(
        "classify_claims",
        route_after_claim_classify,
        {
            "translate_claims": "translate_claims",
            "translate_claims_indep": "translate_claims_indep",
            "translate_claims_dep": "translate_claims_dep",
        },
    )

    g.add_edge("translate_default", "route_section")
    g.add_edge("translate_claims", "postprocess_claims")
    g.add_edge("translate_claims_indep", "postprocess_claims")
    g.add_edge("translate_claims_dep", "postprocess_claims")
    g.add_edge("postprocess_claims", "route_section")
    g.add_edge("translate_abstract", "route_section")

    return g.compile()


# ============================================================
# 7) Apply translations back into DOCX
# ============================================================

def apply_translations_to_docx(
    src_docx_path: str | Path,
    translations: Dict[str, str],
    out_docx_path: str | Path,
) -> None:
    doc = Document(str(src_docx_path))

    for block_id, translated in translations.items():
        if block_id.startswith("p:"):
            _, idx = block_id.split(":", 1)
            pi = int(idx)
            if 0 <= pi < len(doc.paragraphs):
                doc.paragraphs[pi].text = translated

        elif block_id.startswith("cell:"):
            parts = block_id.split(":")
            if len(parts) == 5:
                _, ti, ri, ci, pi = parts
                ti = int(ti)
                ri = int(ri)
                ci = int(ci)
                pi = int(pi)

                if 0 <= ti < len(doc.tables):
                    table = doc.tables[ti]
                    if 0 <= ri < len(table.rows):
                        row = table.rows[ri]
                        if 0 <= ci < len(row.cells):
                            cell = row.cells[ci]
                            if 0 <= pi < len(cell.paragraphs):
                                cell.paragraphs[pi].text = translated

    doc.save(str(out_docx_path))


# ============================================================
# 8) Comparison report
# ============================================================


def _escape_markdown_cell(text: str) -> str:
    return (text or "").replace("\n", "<br>").replace("|", r"\|")


def _escape_tsv_cell(text: str) -> str:
    return (text or "").replace("\t", " ").replace("\n", "\\n")


def write_comparison_reports(
    *,
    src_docx_path: str | Path,
    out_docx_path: str | Path,
    chunk_spans: List[tuple[List[Block], int, int]],
) -> tuple[Path, Path]:
    """
    Write comparison reports for debugging/review:
    - chunk-level markdown: source vs translated per chunk
    - line-level tsv: one row per block
    """
    src_docx_path = Path(src_docx_path)
    out_docx_path = Path(out_docx_path)

    source_blocks = iter_blocks(src_docx_path)
    output_blocks = iter_blocks(out_docx_path)
    source_by_id = {b.id: (b.text or "") for b in source_blocks}
    output_by_id = {b.id: (b.text or "") for b in output_blocks}
    source_ids = set(source_by_id.keys())
    output_ids = set(output_by_id.keys())
    missing_in_output = sorted(source_ids - output_ids)
    extra_in_output = sorted(output_ids - source_ids)

    compare_dir = out_docx_path.parent / "comparisons"
    compare_dir.mkdir(parents=True, exist_ok=True)

    base = out_docx_path.stem
    chunk_md_path = compare_dir / f"{base}.compare.chunks.md"
    line_tsv_path = compare_dir / f"{base}.compare.lines.tsv"

    created_at = datetime.now().isoformat(timespec="seconds")

    md_lines: List[str] = [
        f"# Comparison Report: {base}",
        "",
        f"- Created: {created_at}",
        f"- Source DOCX: {src_docx_path}",
        f"- Output DOCX: {out_docx_path}",
        f"- Source blocks: {len(source_blocks)}",
        f"- Output blocks: {len(output_blocks)}",
        f"- Total chunks: {len(chunk_spans)}",
        f"- Missing blocks in output: {len(missing_in_output)}",
        f"- Extra blocks in output: {len(extra_in_output)}",
        "",
    ]

    tsv_lines: List[str] = [
        "chunk_idx\tline_in_chunk\tblock_id\tmatch_status\tsource_text\ttranslated_text"
    ]

    for chunk_idx, (chunk, s, e) in enumerate(chunk_spans, start=1):
        src_lines = [source_by_id.get(b.id, "") for b in chunk]
        tgt_lines = [output_by_id.get(b.id, "") for b in chunk]

        src_chunk_text = "\n".join(src_lines).strip()
        tgt_chunk_text = "\n".join(tgt_lines).strip()

        md_lines.extend(
            [
                f"## Chunk {chunk_idx} (blocks[{s}:{e}], lines={len(chunk)})",
                "",
                "### Source",
                "```text",
                src_chunk_text,
                "```",
                "",
                "### Translation",
                "```text",
                tgt_chunk_text,
                "```",
                "",
                "| Block ID | Source | Translation |",
                "| --- | --- | --- |",
            ]
        )

        for line_in_chunk, b in enumerate(chunk, start=1):
            source_text = source_by_id.get(b.id, "")
            translated_text = output_by_id.get(b.id, "")
            match_status = "matched" if b.id in output_by_id else "missing_in_output"
            md_lines.append(
                f"| {b.id} | {_escape_markdown_cell(source_text)} | {_escape_markdown_cell(translated_text)} |"
            )

            tsv_lines.append(
                "\t".join(
                    [
                        str(chunk_idx),
                        str(line_in_chunk),
                        b.id,
                        match_status,
                        _escape_tsv_cell(source_text),
                        _escape_tsv_cell(translated_text),
                    ]
                )
            )

        md_lines.append("")

    if missing_in_output:
        md_lines.extend(
            [
                "## Unmatched: Missing In Output",
                "",
                "| Block ID | Source |",
                "| --- | --- |",
            ]
        )
        for bid in missing_in_output:
            md_lines.append(f"| {bid} | {_escape_markdown_cell(source_by_id.get(bid, ''))} |")
            tsv_lines.append(
                "\t".join(
                    [
                        "NA",
                        "NA",
                        bid,
                        "missing_in_output",
                        _escape_tsv_cell(source_by_id.get(bid, "")),
                        "",
                    ]
                )
            )
        md_lines.append("")

    if extra_in_output:
        md_lines.extend(
            [
                "## Unmatched: Extra In Output",
                "",
                "| Block ID | Output |",
                "| --- | --- |",
            ]
        )
        for bid in extra_in_output:
            md_lines.append(f"| {bid} | {_escape_markdown_cell(output_by_id.get(bid, ''))} |")
            tsv_lines.append(
                "\t".join(
                    [
                        "NA",
                        "NA",
                        bid,
                        "extra_in_output",
                        "",
                        _escape_tsv_cell(output_by_id.get(bid, "")),
                    ]
                )
            )
        md_lines.append("")

    chunk_md_path.write_text("\n".join(md_lines), encoding="utf-8")
    line_tsv_path.write_text("\n".join(tsv_lines) + "\n", encoding="utf-8")
    return chunk_md_path, line_tsv_path


# ============================================================
# 9) End-to-end runner
# ============================================================


def translate_docx(
    src_docx_path: str | Path,
    out_docx_path: str | Path,
    *,
    model: str,
    target_lang: str = "English",
    max_chars_per_chunk: int = 2500,
    context_chars: int = 1000,
    chunk_overlap: int = 0,
    base_url: Optional[str] = None,
    api_key: Optional[str] = None,
    prompt_default: str = "patent_kr2en_body_v1",
    prompt_claims: str = "patent_kr2en_claims_v1",
    prompt_claims_indep: str = "patent_kr2en_claims_indep_v1",
    prompt_claims_dep: str = "patent_kr2en_claims_dep_v1",
    prompt_abstract: str = "patent_kr2en_abstract_v1",
    compare: bool = False,
    # NEW: "whole-section" mode with adaptive fallback
    max_section_chars: int = 80_000,
) -> None:
    print("Start chunking...")
    blocks_list = list(iter_blocks(src_docx_path))

    # NEW: Section-sized chunks first; fallback split if a section is too large
    chunk_spans = build_section_chunks_with_fallback(
        blocks_list,
        max_section_chars=max_section_chars,
        max_chars_per_chunk=max_chars_per_chunk,
        chunk_overlap=chunk_overlap,
    )

    chunks = [c for (c, _, _) in chunk_spans]

    # IMPORTANT: build_contexts needs a stable Sequence; use blocks_list (not an iterator)
    contexts_raw = build_contexts(
        blocks_list,
        [(s, e) for (_, s, e) in chunk_spans],
        context_chars=context_chars,
    )
    contexts = [{"before": b, "after": a} for (b, a) in contexts_raw]
    print(f"Finish chunking: {len(chunks)} chunks (section-sized; fallback if > {max_section_chars} chars)")

    print("Build graph...")
    app = build_translation_graph()

    print("Start invoking...")
    final = app.invoke(
        {
            "chunks": chunks,
            "contexts": contexts,
            "i": 0,
            "results": {},
            "model": model,
            "base_url": base_url,
            "api_key": api_key,
            "target_lang": target_lang,
            "section": "default",
            "prompt_default": prompt_default,
            "prompt_claims": prompt_claims,
            "prompt_claims_indep": prompt_claims_indep,
            "prompt_claims_dep": prompt_claims_dep,
            "prompt_abstract": prompt_abstract,
            "abstract_word_count": 0,
            "last_abstract_block_id": None,
            "abstract_completed": False,
            "glossary": {},
            "prev_translated_text": "",
            "claim_preambles": {},
            "id_mismatch_detected": False,
            "id_mismatch_chunks": [],
            "empty_fallback_count": 0,
            "empty_fallback_chunks": [],
        },
        # Each chunk takes up to ~5 graph steps (route→classify→translate→postprocess→route).
        # Add a fixed buffer of 50 to cover the final finalize_abstract pass.
        config={"recursion_limit": max(100, len(chunks) * 6 + 50)},
    )

    # If the document ends while still in abstract (or word count exists), finalize again defensively
    if final.get("abstract_word_count", 0) or final.get("section") == "abstract":
        final = node_route_section({**final, "i": len(chunks)})

    if final.get("id_mismatch_detected"):
        details = "; ".join(final.get("id_mismatch_chunks", []))
        print(f"[WARN] ID mismatch detected after retries (partial source text preserved): {details}")

    fallback_total = int(final.get("empty_fallback_count", 0) or 0)
    if fallback_total > 0:
        details = "; ".join(final.get("empty_fallback_chunks", []))
        print(f"[QUALITY] empty-output fallbacks applied: {fallback_total} ({details})")
    else:
        print("[QUALITY] empty-output fallbacks applied: 0")

    print("Apply translations...")
    apply_translations_to_docx(src_docx_path, final["results"], out_docx_path)

    if compare:
        chunk_md_path, line_tsv_path = write_comparison_reports(
            src_docx_path=src_docx_path,
            out_docx_path=out_docx_path,
            chunk_spans=chunk_spans,
        )
        print(f"[COMPARE] Chunk report: {chunk_md_path}")
        print(f"[COMPARE] Line report: {line_tsv_path}")

    print("Done.")

